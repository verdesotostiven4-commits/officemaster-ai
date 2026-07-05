from pathlib import Path
from docx import Document


def create_word_document(prompt: str, title: str, style: str, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    file_path = output_dir / 'officemaster_word.docx'

    document = Document()
    document.add_heading(title or 'Documento generado por OfficeMaster AI', level=0)
    document.add_paragraph('Estilo seleccionado: ' + style)
    document.add_paragraph('Solicitud del usuario:')
    document.add_paragraph(prompt)

    document.add_heading('Introduccion', level=1)
    document.add_paragraph('Este documento es una base inicial generada por OfficeMaster AI.')

    document.add_heading('Desarrollo', level=1)
    document.add_paragraph('Aqui se organizara la informacion principal con estructura clara y profesional.')

    document.add_heading('Conclusion', level=1)
    document.add_paragraph('La conclusion resumira las ideas principales del documento.')

    document.add_heading('Referencias', level=1)
    document.add_paragraph('En una fase futura se podra generar esta seccion con formato APA 7.')

    document.save(file_path)
    return file_path
